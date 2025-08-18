

import os, json, math, argparse, time, random
from collections import Counter
from typing import List, Dict, Any

import numpy as np
import pandas as pd

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from gensim.models import FastText
import faiss

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

NUM_TEST       = 50_000
NUM_RUNS       = 1
TOP_K          = 100
SEED           = 42
OUTPUT_DIR     = "dl_output_fasttext_eval"

FT_VECTOR_SIZE = 100
FT_WINDOW      = 5
FT_MIN_COUNT   = 3
FT_SG          = 1    
FT_EPOCHS      = 5

def set_seed(seed: int):
    random.seed(seed); np.random.seed(seed)

def read_jsonl(path: str) -> pd.DataFrame:
    df = pd.read_json(path, lines=True)
    return df

def process_categories_for_df(categories):
    if categories and isinstance(categories, list):
        filtered = [cat[-1] for cat in categories if cat and cat[0] != "Clothing, Shoes & Jewelry"]
        return filtered[-1] if filtered else categories[-1][-1]
    return ""

def process_salesRank(salesRank):
    if salesRank and isinstance(salesRank, dict):
        return list(salesRank.keys())[0]
    return ""

def preprocess_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df["title"]      = df.get("title", "").fillna("").astype(str)
    df["brand"]      = df.get("brand", "").fillna("").astype(str)
    df["categories"] = df.get("categories", []).apply(process_categories_for_df)
    df["salesRank"]  = df.get("salesRank", {}).apply(process_salesRank)
    df["combined_text"] = (df["title"] + " " + df["categories"] + " " + df["salesRank"] + " " + df["brand"]).str.strip()
    return df

def build_asin_to_categories_from_file(path: str) -> Dict[str, Any]:
    """asin -> categories (list of lists) from raw JSONL."""
    asin_to_cat = {}
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            try:
                obj = json.loads(line)
                asin_to_cat[obj["asin"]] = obj.get("categories", [])
            except Exception:
                continue
    return asin_to_cat

def extract_main_category_list(categories_ll):
    if not categories_ll or not isinstance(categories_ll, list):
        return []
    best, best_len = [], 0
    for sub in categories_ll:
        if not sub: continue
        filt = [c for c in sub if c != "Clothing, Shoes & Jewelry"]
        if len(filt) > best_len:
            best, best_len = filt, len(filt)
    return best

def main_category_name(categories_ll) -> str:
    best = extract_main_category_list(categories_ll)
    return best[-1] if best else ""

def build_main_category_index(df_vec: pd.DataFrame, asin_to_cat: Dict[str, Any]):
    cat2idxs = {}
    main_cats = []
    for i, row in df_vec.iterrows():
        a = row["asin"]
        mc = main_category_name(asin_to_cat.get(a, []))
        main_cats.append(mc)
        if mc:
            cat2idxs.setdefault(mc, []).append(i)
    return cat2idxs, main_cats

def stratified_test_sampling(cat2idxs: Dict[str, List[int]], total_queries: int, seed: int = 42) -> List[int]:
    rng = np.random.default_rng(seed)
    cats = list(cat2idxs.keys())
    C = len(cats)
    if C == 0:
        raise ValueError("Không có main category để cân bằng test.")
    base = total_queries // C
    rem  = total_queries % C
    targets = {c: base for c in cats}
    if rem > 0:
        for c in rng.choice(cats, size=rem, replace=False):
            targets[c] += 1
    chosen, deficit, surplus = [], 0, []
    for c in cats:
        pool = cat2idxs[c]
        take = min(len(pool), targets[c])
        if take > 0:
            chosen.extend(rng.choice(pool, size=take, replace=False).tolist())
        if len(pool) < targets[c]:
            deficit += (targets[c] - len(pool))
        elif len(pool) > targets[c]:
            surplus.append(c)
    if deficit > 0 and surplus:
        extra_pool = []
        for c in surplus:
            pool = cat2idxs[c]
            extra = len(pool) - targets[c]
            if extra > 0:
                cand = rng.choice(pool, size=min(extra*2, len(pool)), replace=False).tolist()
                extra_pool.extend(cand)
        chosen_set = set(chosen)
        extra_pool = [x for x in extra_pool if x not in chosen_set]
        chosen.extend(extra_pool[:min(deficit, len(extra_pool))])
    if len(chosen) > total_queries:
        chosen = rng.choice(chosen, size=total_queries, replace=False).tolist()
    rng.shuffle(chosen)
    return chosen

def train_fasttext(tokenized: List[List[str]],
                   vector_size=FT_VECTOR_SIZE,
                   window=FT_WINDOW,
                   min_count=FT_MIN_COUNT,
                   sg=FT_SG,
                   epochs=FT_EPOCHS,
                   seed=SEED) -> FastText:
    random.seed(seed); np.random.seed(seed)
    model = FastText(vector_size=vector_size, window=window, min_count=min_count,
                     workers=os.cpu_count() or 4, sg=sg)
    if not tokenized or all(len(s)==0 for s in tokenized):
        raise ValueError("Tokenized corpus is empty; cannot train FastText.")
    model.build_vocab(tokenized)
    model.train(tokenized, total_examples=len(tokenized), epochs=epochs)
    return model

def build_sentence_vectors(model: FastText, tokenized: List[List[str]]) -> List[np.ndarray | None]:
    vecs = []
    for tokens in tokenized:
        arrs = [model.wv[w] for w in tokens if w in model.wv]
        if not arrs:
            vecs.append(None)
        else:
            v = np.mean(arrs, axis=0, dtype=np.float32)
            n = np.linalg.norm(v)
            vecs.append((v / n).astype(np.float32) if n > 0 else None)
    return vecs

def build_faiss(vectors: List[np.ndarray | None], index_factory: str = "Flat"):
    keep_idx = [i for i, v in enumerate(vectors) if v is not None]
    mat = np.vstack([vectors[i] for i in keep_idx]).astype(np.float32)
    dim = mat.shape[1]

    if index_factory.lower() in ("flat", "flatip", ""):
        index = faiss.IndexFlatIP(dim)
        index.add(mat)
        desc = "IndexFlatIP"
    else:
        index = faiss.index_factory(dim, index_factory, faiss.METRIC_INNER_PRODUCT)
        if not index.is_trained:
            index.train(mat)
        index.add(mat)
        desc = f"ANN({index_factory}, metric=IP)"
    return index, keep_idx, mat, desc

def search_drop_self(index, q_vec: np.ndarray, q_idx: int, topk: int):
    D, I = index.search(np.array([q_vec], dtype=np.float32), topk + 1)
    if I[0][0] == q_idx:
        return I[0][1:topk+1], D[0][1:topk+1]
    return I[0][:topk], D[0][:topk]

def categories_match(cat1, cat2) -> bool:
    """Relevant nếu giao nhau >= 2 categories (bỏ gốc 'Clothing, Shoes & Jewelry')."""
    def to_set(lst):
        s = set()
        for sub in lst:
            for c in sub:
                if c != "Clothing, Shoes & Jewelry":
                    s.add(c)
        return s
    s1, s2 = to_set(cat1), to_set(cat2)
    return len(s1 & s2) >= 2

def dcg(rels, k):   return sum(((2**r - 1) / math.log2(i + 2)) for i, r in enumerate(rels[:k]))
def idcg(k):        return dcg([1]*k, k)

def evaluate(df_vec: pd.DataFrame, vectors_mat: np.ndarray, test_idx: List[int],
             asin_to_cat: Dict[str, Any], index, topk: int) -> Dict[str, float]:
    P_at = {5:0.0, 10:0.0, 20:0.0}
    ndcg = {5:0.0, 10:0.0, 20:0.0}
    mrr  = 0.0
    correct = 0
    total   = 0
    N = len(test_idx)

    for t, idx in enumerate(test_idx, 1):
        if t % 1000 == 0:
            print(f"  test {t}/{N}")
        q_vec = vectors_mat[idx]
        q_asin = df_vec.iloc[idx]["asin"]
        I, D = search_drop_self(index, q_vec, q_idx=idx, topk=topk)

        rels = []
        first = None
        q_cat = asin_to_cat.get(q_asin, [])
        for rank, ridx in enumerate(I):
            r_asin = df_vec.iloc[ridx]["asin"]
            ok = categories_match(q_cat, asin_to_cat.get(r_asin, []))
            rels.append(1 if ok else 0)
            if ok and first is None:
                first = rank + 1

        for k in P_at:
            if len(rels) >= k:
                P_at[k] += sum(rels[:k]) / k
                ndcg[k] += dcg(rels, k) / idcg(k)
        if first:
            mrr += 1.0 / first
        correct += sum(rels)
        total   += len(rels)

    metrics = {
        "Precision@5":  P_at[5] / max(1, N),
        "Precision@10": P_at[10] / max(1, N),
        "Precision@20": P_at[20] / max(1, N),
        "MRR":          mrr / max(1, N),
        "NDCG@5":       ndcg[5] / max(1, N),
        "NDCG@10":      ndcg[10] / max(1, N),
        "NDCG@20":      ndcg[20] / max(1, N),
        "Accuracy":     (correct / total) if total else 0.0
    }
    return metrics

def plot_metrics(metrics: dict, out_path: str, title: str):
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    labels, vals = list(metrics.keys()), list(metrics.values())
    plt.figure(figsize=(12,6))
    bars = plt.bar(labels, vals)
    plt.ylim(0,1); plt.xticks(rotation=45); plt.title(title)
    for b in bars:
        plt.text(b.get_x()+b.get_width()/2, b.get_height()+0.01, f"{b.get_height():.4f}",
                 ha="center", va="bottom")
    plt.tight_layout(); plt.savefig(out_path); plt.close()

def plot_category_distribution(cat_counts: Counter, out_path: str):
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    cats, cnts = list(cat_counts.keys()), list(cat_counts.values())
    plt.figure(figsize=(12, max(6, min(18, len(cats)//5 + 4))))
    plt.barh(cats, cnts)
    plt.xlabel("Số lượng sản phẩm"); plt.ylabel("Danh mục (main)")
    plt.title("Phân bố main category (trong df_vec)")
    plt.tight_layout(); plt.savefig(out_path); plt.close()

def export_docx_report(avg_metrics: dict, cat_counts: Counter, out_docx: str,
                       metrics_img: str, dist_img: str):
    if not HAVE_DOCX:
        print("⚠️ python-docx chưa cài; bỏ qua DOCX.")
        return
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    title = doc.add_heading("Kết quả đánh giá FastText (Content-based)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Phân bố main category trong tập dùng để search (df_vec).")
    doc.add_picture(dist_img)

    doc.add_paragraph("Trung bình các chỉ số qua các lần chạy.")
    doc.add_picture(metrics_img)

    table = doc.add_table(rows=len(avg_metrics)+1, cols=2)
    table.style = "Table Grid"
    table.cell(0,0).text = "Chỉ số"
    table.cell(0,1).text = "Giá trị"
    i = 1
    for k, v in avg_metrics.items():
        table.cell(i,0).text = k
        table.cell(i,1).text = f"{v:.4f}"
        i += 1
    doc.save(out_docx)
    print(f"✅ Saved DOCX: {out_docx}")

def export_recs_for_test(df_vec: pd.DataFrame, vectors_mat: np.ndarray,
                         test_idx: List[int], index, topk: int, out_csv: str):
    rows = []
    for t, idx in enumerate(test_idx, 1):
        if t % 2000 == 0:
            print(f"  export rec {t}/{len(test_idx)}")
        q_vec = vectors_mat[idx]
        I, D = search_drop_self(index, q_vec, q_idx=idx, topk=topk)
        rec_asins = [df_vec.iloc[j]["asin"] for j in I]
        rows.append({
            "asin": df_vec.iloc[idx]["asin"],
            "recommend_asins": ",".join(rec_asins),
            "avg_similarity": float(np.mean(D)) if len(D) else 0.0
        })
    pd.DataFrame(rows).to_csv(out_csv, index=False)
    print(f"✅ Saved recommendations: {out_csv}")

def main(args):
    t0 = time.time()
    set_seed(args.seed)
    os.makedirs(args.out, exist_ok=True)

    print("🔹 Read JSONL...")
    df_raw = read_jsonl(args.jsonl)
    if len(df_raw) == 0:
        raise ValueError("No rows read from JSONL. Kiểm tra đường dẫn file --jsonl.")
    print(f"Rows read: {len(df_raw)}")

    print("🔹 Preprocess...")
    df = preprocess_df(df_raw)

    print("🔹 Tokenize...")
    tokenized = [str(x).lower().split() for x in df["combined_text"].tolist()]

    print(f"🔹 Train FastText (dim={args.ft_dim}, win={args.ft_window}, minc={args.ft_min_count}, "
          f"sg={args.ft_sg}, epochs={args.ft_epochs})...")
    ft_model = train_fasttext(
        tokenized,
        vector_size=args.ft_dim,
        window=args.ft_window,
        min_count=args.ft_min_count,
        sg=args.ft_sg,
        epochs=args.ft_epochs,
        seed=args.seed
    )

    print("🔹 Build sentence vectors + L2 norm...")
    vectors = build_sentence_vectors(ft_model, tokenized)

    print(f"🔹 Build FAISS index ({args.index_factory})...")
    index, keep_idx, mat, desc = build_faiss(vectors, index_factory=args.index_factory)
    df_vec = df.iloc[keep_idx].reset_index(drop=True)
    print(f"    → {desc}, ntotal={index.ntotal}, dim={mat.shape[1]}, df_vec={len(df_vec)}")

    print("🔹 asin->categories mapping (from file)...")
    asin_to_cat = build_asin_to_categories_from_file(args.jsonl)

    print("🔹 Build cat2idxs for stratified test...")
    cat2idxs, main_cats = build_main_category_index(df_vec, asin_to_cat)
    cat_counts = Counter([c for c in main_cats if c])
    print("Main categories:", len(cat_counts))
    print("Top-10:", cat_counts.most_common(10))

    if args.num_test > len(df_vec):
        raise ValueError(f"NUM_TEST ({args.num_test}) > available rows ({len(df_vec)}). Reduce num_test or increase data.")

    all_metrics = []
    for run in range(1, args.num_runs + 1):
        print(f"\n=== RUN {run}/{args.num_runs} ===")
        test_idx = stratified_test_sampling(cat2idxs, total_queries=args.num_test, seed=args.seed + run)
        m = evaluate(df_vec, mat, test_idx, asin_to_cat, index, topk=args.topk)
        all_metrics.append(m)
        for k, v in m.items(): print(f"{k}: {v:.4f}")
        plot_metrics(m, os.path.join(args.out, f"metrics_run_{run}.png"), f"Run {run}")
        if args.export_recs:
            out_csv = os.path.join(args.out, f"test_recs_run_{run}.csv")
            export_recs_for_test(df_vec, mat, test_idx, index, topk=args.topk, out_csv=out_csv)

    avg = {k: float(np.mean([mm[k] for mm in all_metrics])) for k in all_metrics[0]}
    print("\n=== AVERAGE OVER RUNS ===")
    for k, v in avg.items(): print(f"{k}: {v:.4f}")

    pd.DataFrame([avg]).to_csv(os.path.join(args.out, "avg_metrics.csv"), index=False)
    dist_img = os.path.join(args.out, "category_distribution.png")
    avg_img  = os.path.join(args.out, "metrics_avg.png")
    plot_category_distribution(cat_counts, dist_img)
    plot_metrics(avg, avg_img, "Average over runs")
    if args.export_docx:
        export_docx_report(avg, cat_counts, os.path.join(args.out, "results_fastText.docx"),
                           metrics_img=avg_img, dist_img=dist_img)

    print(f"\n✅ Done in {(time.time()-t0)/60:.1f} min. Output: {args.out}")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--jsonl", required=True, help="Path to meta_Clothing_Shoes_and_Jewelry*.jsonl")
    ap.add_argument("--out", default=OUTPUT_DIR)
    ap.add_argument("--ft_dim", type=int, default=FT_VECTOR_SIZE)
    ap.add_argument("--ft_window", type=int, default=FT_WINDOW)
    ap.add_argument("--ft_min_count", type=int, default=FT_MIN_COUNT)
    ap.add_argument("--ft_sg", type=int, default=FT_SG)
    ap.add_argument("--ft_epochs", type=int, default=FT_EPOCHS)
    ap.add_argument("--num_test", type=int, default=NUM_TEST)
    ap.add_argument("--num_runs", type=int, default=NUM_RUNS)
    ap.add_argument("--topk", type=int, default=TOP_K)
    ap.add_argument("--seed", type=int, default=SEED)
    ap.add_argument("--index_factory", default="HNSW32", help="Flat | HNSW32 | IVF4096,Flat | IVF8192,PQ64 ...")
    ap.add_argument("--export_recs", action="store_true", help="Export CSV recs for test queries")
    ap.add_argument("--export_docx", action="store_true", help="Export DOCX report")
    args = ap.parse_args()
    main(args)
